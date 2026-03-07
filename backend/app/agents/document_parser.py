"""
Document Parser Agent - Extracts text from PDF/DOCX files with OCR support
"""
import logging
from typing import Dict, Any, Optional
import pdfplumber
from docx import Document
import io
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    convert_from_bytes = None
    Image = None

from app.services.storage_factory import storage as storage_service

logger = logging.getLogger(__name__)


class DocumentParserAgent:
    """Agent for parsing document files and extracting raw text"""
    
    def __init__(self):
        self.storage_service = storage_service
    
    async def parse_document(self, blob_url: str) -> Dict[str, Any]:
        """
        Parse a document and extract text
        
        Args:
            blob_url: URL of the blob storage file
            
        Returns:
            Dictionary containing parsed text and metadata
        """
        logger.info(f"Parsing document from: {blob_url}")
        
        try:
            # Download file
            file_content = await self.storage_service.download_file(blob_url)
            
            # Determine file type from URL
            if blob_url.endswith('.pdf') or 'pdf' in blob_url.lower():
                result = await self._parse_pdf(file_content)
            elif blob_url.endswith('.docx') or 'docx' in blob_url.lower():
                result = await self._parse_docx(file_content)
            elif blob_url.endswith('.txt') or 'text/plain' in blob_url.lower():
                result = await self._parse_txt(file_content)
            elif any(blob_url.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp']):
                result = await self._parse_image(file_content, blob_url)
            else:
                raise ValueError(f"Unsupported file type: {blob_url}")
            
            # Log with appropriate metric
            if result.get('page_count'):
                logger.info(f"Parsed document: {result['page_count']} pages, {len(result['full_text'])} characters")
            elif result.get('paragraph_count'):
                logger.info(f"Parsed document: {result['paragraph_count']} paragraphs, {len(result['full_text'])} characters")
            else:
                logger.info(f"Parsed document: {len(result['full_text'])} characters")
            
            return result
            
        except Exception as e:
            logger.error(f"Error parsing document: {e}", exc_info=True)
            raise
    
    async def _parse_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Parse PDF file using pdfplumber"""
        try:
            pdf_file = io.BytesIO(file_content)
            
            pages = []
            full_text = ""
            total_pages = 0
            
            with pdfplumber.open(pdf_file) as pdf:
                total_pages = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text
                    text = page.extract_text() or ""
                    
                    # Also extract tables if present
                    tables = page.extract_tables()
                    tables_text = ""
                    if tables:
                        for table in tables:
                            # Convert table to text
                            table_str = "\n".join([
                                " | ".join([str(cell) if cell else "" for cell in row])
                                for row in table
                            ])
                            tables_text += "\n\n" + table_str + "\n\n"
                    
                    page_text = text + tables_text
                    
                    pages.append({
                        "page_number": page_num,
                        "text": page_text
                    })
                    full_text += f"\n\n--- Page {page_num} ---\n\n{page_text}"
            
            # Check if PDF is scanned (minimal text extracted)
            avg_chars_per_page = len(full_text.strip()) / max(total_pages, 1)
            if avg_chars_per_page < 100 and OCR_AVAILABLE:
                logger.info(f"Detected scanned PDF (avg {avg_chars_per_page:.1f} chars/page), applying OCR...")
                return await self._parse_pdf_with_ocr(file_content)
            
            return {
                "full_text": full_text.strip(),
                "pages": pages,
                "page_count": total_pages,
                "file_type": "pdf",
                "metadata": {
                    "parser": "pdfplumber"
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise
    
    async def _parse_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Parse DOCX file with image OCR support"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            paragraphs = []
            full_text = ""
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        "index": i,
                        "text": para.text
                    })
                    full_text += para.text + "\n\n"
            
            # Also extract tables
            tables_text = []
            for table in doc.tables:
                table_text = "\n".join([
                    " | ".join([cell.text for cell in row.cells])
                    for row in table.rows
                ])
                tables_text.append(table_text)
                full_text += "\n\n" + table_text + "\n\n"
            
            # Extract images and apply OCR if available
            images_text = []
            image_count = 0
            
            if OCR_AVAILABLE:
                try:
                    # Extract images from document relationships
                    for rel in doc.part.rels.values():
                        if "image" in rel.target_ref:
                            try:
                                image_data = rel.target_part.blob
                                image = Image.open(io.BytesIO(image_data))
                                
                                # Apply OCR to image
                                ocr_text = pytesseract.image_to_string(image, lang='eng')
                                
                                if ocr_text.strip():
                                    images_text.append(ocr_text.strip())
                                    full_text += f"\n\n--- Image {image_count + 1} (OCR) ---\n\n{ocr_text.strip()}\n\n"
                                    image_count += 1
                                    logger.info(f"Extracted {len(ocr_text)} characters from embedded image {image_count}")
                            except Exception as img_err:
                                logger.warning(f"Failed to extract image {image_count + 1}: {img_err}")
                                continue
                except Exception as ocr_err:
                    logger.warning(f"Image extraction failed: {ocr_err}")
            
            # Check if document is mostly images (minimal text extracted)
            text_from_docx = len(full_text.strip()) - sum(len(img) for img in images_text)
            has_images = image_count > 0
            mostly_images = has_images and text_from_docx < 100
            
            metadata = {
                "image_count": image_count,
                "has_ocr": has_images and OCR_AVAILABLE,
            }
            
            if mostly_images:
                metadata["note"] = "Document contains primarily images with OCR-extracted text"
            
            return {
                "full_text": full_text.strip(),
                "paragraphs": paragraphs,
                "tables": tables_text,
                "images_text": images_text,
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "image_count": image_count,
                "file_type": "docx",
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise
    
    async def _parse_txt(self, file_content: bytes) -> Dict[str, Any]:
        """Parse plain text file"""
        try:
            # Try common encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # If all encodings fail, use utf-8 with error handling
                text = file_content.decode('utf-8', errors='replace')
            
            # Split into paragraphs (preserve structure)
            paragraphs = []
            for i, para in enumerate(text.split('\n\n')):
                if para.strip():
                    paragraphs.append({
                        "index": i,
                        "text": para.strip()
                    })
            
            return {
                "full_text": text.strip(),
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "file_type": "txt",
                "metadata": {
                    "parser": "text_decoder"
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing TXT: {e}")
            raise
    
    async def _parse_pdf_with_ocr(self, file_content: bytes) -> Dict[str, Any]:
        """Parse scanned PDF using OCR"""
        if not OCR_AVAILABLE:
            raise RuntimeError("OCR libraries not available. Install: pip install pytesseract pdf2image pillow")
        
        try:
            # Convert PDF pages to images
            images = convert_from_bytes(file_content, dpi=300)
            
            pages = []
            full_text = ""
            
            for page_num, image in enumerate(images, start=1):
                # Apply OCR to extract text
                text = pytesseract.image_to_string(image, lang='eng')
                
                pages.append({
                    "page_number": page_num,
                    "text": text
                })
                full_text += f"\n\n--- Page {page_num} ---\n\n{text}"
            
            logger.info(f"OCR extracted {len(full_text)} characters from {len(images)} pages")
            
            return {
                "full_text": full_text.strip(),
                "pages": pages,
                "page_count": len(images),
                "file_type": "pdf",
                "metadata": {
                    "parser": "pytesseract_ocr",
                    "ocr_applied": True
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF with OCR: {e}")
            raise
    
    async def _parse_image(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse image file using OCR"""
        if not OCR_AVAILABLE:
            raise RuntimeError("OCR libraries not available. Install: pip install pytesseract pillow")
        
        try:
            # Open image
            image = Image.open(io.BytesIO(file_content))
            
            # Apply OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            logger.info(f"OCR extracted {len(text)} characters from image")
            
            return {
                "full_text": text.strip(),
                "pages": [{
                    "page_number": 1,
                    "text": text
                }],
                "page_count": 1,
                "file_type": "image",
                "metadata": {
                    "parser": "pytesseract_ocr",
                    "ocr_applied": True,
                    "image_format": filename.split('.')[-1].upper()
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing image with OCR: {e}")
            raise
    
    def chunk_text(
        self,
        text: str,
        max_chunk_size: int = 1000,
        overlap: int = 200,
    ) -> list:
        """
        Split text into overlapping chunks for downstream AI processing.

        Each chunk is represented as a dict so callers can store additional
        metadata alongside the text without changing this method's signature.

        Args:
            text: Full document text.
            max_chunk_size: Maximum characters per chunk (default 1000).
            overlap: Characters shared between adjacent chunks (default 200).

        Returns:
            List of dicts: [{"chunk_index": int, "chunk_text": str}, ...]
        """
        if not text:
            return []

        if max_chunk_size <= overlap:
            raise ValueError("max_chunk_size must be greater than overlap")

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + max_chunk_size
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({"chunk_index": chunk_index, "chunk_text": chunk_text})
                chunk_index += 1
            # Advance by (max_chunk_size - overlap) so consecutive chunks share `overlap` chars
            start += max_chunk_size - overlap

        logger.info(f"Created {len(chunks)} chunks from {len(text)} characters")
        return chunks

    def extract_paragraphs(self, parse_result: dict) -> list:
        """
        Normalise paragraph output from both _parse_pdf and _parse_docx into a
        flat list of non-empty paragraph strings.

        Args:
            parse_result: Return value of parse_document().

        Returns:
            Ordered list of non-empty paragraph strings.
        """
        paragraphs: list = []

        if parse_result.get("file_type") == "pdf":
            for page in parse_result.get("pages", []):
                for line in page.get("text", "").splitlines():
                    stripped = line.strip()
                    if stripped:
                        paragraphs.append(stripped)
        else:
            # DOCX
            for item in parse_result.get("paragraphs", []):
                text = item.get("text", "").strip() if isinstance(item, dict) else str(item).strip()
                if text:
                    paragraphs.append(text)

        return paragraphs
