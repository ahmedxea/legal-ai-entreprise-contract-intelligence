"""
Test DOCX files with embedded images (OCR extraction)
"""
import sys
import asyncio
from pathlib import Path
sys.path.insert(0, '.')

from io import BytesIO
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from app.agents.document_parser import DocumentParserAgent, OCR_AVAILABLE


async def main():
    print("🧪 Testing DOCX with Embedded Images (OCR)")
    print("=" * 60)
    
    if not OCR_AVAILABLE:
        print("❌ OCR libraries not available!")
        print("   Install with: pip install pytesseract pillow")
        return
    
    print("✅ OCR libraries available")
    
    # Create a test DOCX with some text and an embedded image
    print("\n1️⃣ Creating test DOCX with text and image...")
    
    # Create test image with contract text
    img = Image.new('RGB', (800, 400), color='white')
    draw = ImageDraw.Draw(img)
    
    contract_text = """
    CONTRACT AGREEMENT
    
    This agreement is made between Party A and Party B
    Effective Date: March 7, 2026
    
    Terms and Conditions:
    1. Payment: $50,000
    2. Duration: 12 months
    3. Governing Law: California
    """
    
    # Draw text on image
    draw.text((50, 50), contract_text, fill='black')
    
    # Save image to memory
    img_bytes = BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    
    # Create DOCX
    doc = Document()
    doc.add_paragraph("Contract Document with Image")
    doc.add_paragraph("See embedded contract below:")
    
    # Add image to DOCX
    doc.add_picture(img_bytes, width=docx.shared.Inches(5))
    
    doc.add_paragraph("End of contract document")
    
    # Save DOCX to bytes
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    print("   ✅ Test DOCX created with embedded contract image")
    
    # Test parsing
    print("\n2️⃣ Parsing DOCX with image OCR...")
    parser = DocumentParserAgent()
    
    try:
        result = await parser._parse_docx(docx_bytes.getvalue())
        
        print(f"   ✅ Parsing successful!")
        print(f"\n   📊 Results:")
        print(f"      - Paragraphs: {result.get('paragraph_count', 0)}")
        print(f"      - Tables: {result.get('table_count', 0)}")
        print(f"      - Images: {result.get('image_count', 0)}")
        print(f"      - Total text: {len(result['full_text'])} characters")
        print(f"      - Has OCR: {result.get('metadata', {}).get('has_ocr', False)}")
        
        if result.get('images_text'):
            print(f"\n   📷 Image OCR Text Preview:")
            for i, img_text in enumerate(result['images_text'], 1):
                preview = img_text[:200].replace('\n', ' ')
                print(f"      Image {i}: {preview}...")
        
        print(f"\n   📄 Full Text Preview (first 300 chars):")
        print(f"      {result['full_text'][:300]}...")
        
        # Check if contract terms were extracted
        full_text_lower = result['full_text'].lower()
        found_terms = []
        if 'party a' in full_text_lower or 'party b' in full_text_lower:
            found_terms.append("✅ Parties detected")
        if 'payment' in full_text_lower or '$50,000' in full_text_lower or '50,000' in full_text_lower:
            found_terms.append("✅ Payment terms detected")
        if 'california' in full_text_lower or 'governing law' in full_text_lower:
            found_terms.append("✅ Governing law detected")
        
        if found_terms:
            print(f"\n   🎯 Contract Terms Extracted:")
            for term in found_terms:
                print(f"      {term}")
        
    except Exception as e:
        print(f"   ❌ Parsing failed: {e}")
        import traceback
        traceback.print_exc()
        return
    
    print("\n" + "=" * 60)
    print("✅ DOCX Image OCR test complete!")
    print("\n💡 Your DOCX file with embedded contract images should now work!")


if __name__ == "__main__":
    # Import docx.shared for image sizing
    try:
        import docx.shared
    except ImportError:
        print("❌ python-docx not installed")
        print("   Install with: pip install python-docx")
        sys.exit(1)
    
    asyncio.run(main())
