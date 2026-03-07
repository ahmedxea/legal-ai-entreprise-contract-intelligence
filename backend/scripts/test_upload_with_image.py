"""
Test uploading DOCX with embedded image to replicate user's scenario
"""
import sys
sys.path.insert(0, '.')

from docx import Document
import docx.shared
from PIL import Image, ImageDraw
from io import BytesIO
import requests
import time

print("🧪 Testing DOCX Upload with Embedded Image")
print("=" * 60)

# Create image with contract text
print("\n1️⃣ Creating contract image...")
img = Image.new('RGB', (800, 600), color='white')
draw = ImageDraw.Draw(img)
contract_text = """CONTRACT AGREEMENT

This is a test contract
with embedded image content

Parties: Party A and Party B
Amount: $10,000 payment
Terms: 12 months"""
draw.text((50, 50), contract_text, fill='black')

# Save image to memory
img_bytes = BytesIO()
img.save(img_bytes, format='PNG')
img_bytes.seek(0)

# Create DOCX with embedded image
print("2️⃣ Creating DOCX with embedded image...")
doc = Document()
doc.add_paragraph('Contract document with embedded image:')
doc.add_picture(img_bytes, width=docx.shared.Inches(5))
doc.add_paragraph('End of contract')

# Save DOCX
docx_bytes = BytesIO()
doc.save(docx_bytes)
file_size = len(docx_bytes.getvalue())
docx_bytes.seek(0)

print(f"   File size: {file_size / 1024:.1f} KB")

# Upload to backend
print("\n3️⃣ Uploading to backend...")
files = {'file': ('OCR Test.docx', docx_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}

try:
    response = requests.post(
        'http://localhost:8000/api/contracts/upload?language=en', 
        files=files, 
        timeout=30
    )
    
    print(f"   Status Code: {response.status_code}")
    
    if response.status_code == 200:
        data = response.json()
        contract_id = data['contract_id']
        print(f"   ✅ Upload successful!")
        print(f"   Contract ID: {contract_id}")
        
        print("\n4️⃣ Waiting for processing...")
        for i in range(10):
            time.sleep(2)
            status_response = requests.get(f'http://localhost:8000/api/contracts/{contract_id}')
            contract = status_response.json()
            status = contract['status']
            print(f"   [{i*2}s] Status: {status}")
            
            if status == 'failed':
                print(f"\n   ❌ FAILED!")
                break
            elif status == 'extracted':
                print(f"\n   ✅ EXTRACTED SUCCESSFULLY!")
                
                # Check document_text table
                print("\n5️⃣ Checking extracted text...")
                import sqlite3
                conn = sqlite3.connect('data/contracts.db')
                cursor = conn.cursor()
                cursor.execute("SELECT LENGTH(raw_text), paragraph_count FROM document_text WHERE document_id = ?", (contract_id,))
                row = cursor.fetchone()
                if row:
                    print(f"   Text length: {row[0]} characters")
                    print(f"   Paragraphs: {row[1]}")
                    
                    if row[0] > 100:
                        print(f"\n   ✅✅ SUCCESS! OCR extracted {row[0]} characters from embedded image!")
                    else:
                        print(f"\n   ⚠️  WARNING: Only {row[0]} characters extracted (expected > 100)")
                conn.close()
                break
    else:
        print(f"   ❌ Upload failed!")
        print(f"   Response: {response.text}")
        
except Exception as e:
    print(f"   ❌ Error: {e}")
    import traceback
    traceback.print_exc()

print("\n" + "=" * 60)
