"""
Phase 1 tests — document ingestion pipeline

Test groups:
  1. DocumentParserAgent.chunk_text  (pure unit tests, no I/O)
  2. DocumentParserAgent.extract_paragraphs  (pure unit tests)
  3. DocumentProcessorService.process  (async unit test with mocks)
  4. API layer — upload validation  (TestClient, no real storage/DB)

Run from backend/:
  pytest tests/test_phase1.py -v
"""
import io
import sys
import uuid
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

# ── Path bootstrap ────────────────────────────────────────────────────────────
_BACKEND = Path(__file__).resolve().parent.parent
if str(_BACKEND) not in sys.path:
    sys.path.insert(0, str(_BACKEND))

# ── Helpers ───────────────────────────────────────────────────────────────────

def _make_docx_bytes() -> bytes:
    """Minimal in-memory DOCX with two paragraphs."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph of the contract.")
    doc.add_paragraph("Second paragraph with important terms.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_minimal_pdf_bytes() -> bytes:
    """
    Hand-crafted valid single-page PDF with one text stream.
    No external PDF library required.
    """
    return (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
        b"/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj\n"
        b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"5 0 obj<</Length 44>>\nstream\n"
        b"BT /F1 12 Tf 100 700 Td (Hello World) Tj ET\n"
        b"endstream\nendobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000266 00000 n \n"
        b"0000000347 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n445\n%%EOF\n"
    )


# ═════════════════════════════════════════════════════════════════════════════
# 1. chunk_text — pure unit tests
# ═════════════════════════════════════════════════════════════════════════════

class TestChunkText:
    """DocumentParserAgent.chunk_text produces correct dicts and boundaries."""

    def setup_method(self):
        from app.agents.document_parser import DocumentParserAgent
        self.parser = DocumentParserAgent()

    def test_returns_list_of_dicts(self):
        chunks = self.parser.chunk_text("A" * 500)
        assert isinstance(chunks, list)
        assert len(chunks) > 0
        for c in chunks:
            assert "chunk_index" in c
            assert "chunk_text" in c

    def test_chunk_indices_are_sequential(self):
        chunks = self.parser.chunk_text("B" * 3000)
        indices = [c["chunk_index"] for c in chunks]
        assert indices == list(range(len(indices)))

    def test_overlap_is_applied(self):
        text = "X" * 1200
        chunks = self.parser.chunk_text(text, max_chunk_size=1000, overlap=200)
        # With 1200 chars, chunk_size=1000, overlap=200:
        # chunk 0: [0, 1000], chunk 1: [800, 1200]  → 2 chunks
        assert len(chunks) == 2
        assert chunks[0]["chunk_text"].startswith(chunks[1]["chunk_text"][:200])

    def test_empty_text_returns_empty_list(self):
        assert self.parser.chunk_text("") == []

    def test_text_shorter_than_chunk_size_gives_one_chunk(self):
        chunks = self.parser.chunk_text("Short text.", max_chunk_size=1000, overlap=200)
        assert len(chunks) == 1
        assert chunks[0]["chunk_index"] == 0

    def test_max_chunk_length_respected(self):
        chunks = self.parser.chunk_text("C" * 5000, max_chunk_size=1000, overlap=0)
        for c in chunks:
            assert len(c["chunk_text"]) <= 1000

    def test_invalid_overlap_raises(self):
        with pytest.raises(ValueError):
            self.parser.chunk_text("text", max_chunk_size=100, overlap=100)


# ═════════════════════════════════════════════════════════════════════════════
# 2. extract_paragraphs — pure unit tests
# ═════════════════════════════════════════════════════════════════════════════

class TestExtractParagraphs:
    """DocumentParserAgent.extract_paragraphs normalises both PDF and DOCX output."""

    def setup_method(self):
        from app.agents.document_parser import DocumentParserAgent
        self.parser = DocumentParserAgent()

    def test_pdf_parse_result(self):
        parse_result = {
            "file_type": "pdf",
            "pages": [
                {"page_number": 1, "text": "Line one\n\nLine two\n\nLine three"},
            ],
            "full_text": "Line one\n\nLine two\n\nLine three",
        }
        paras = self.parser.extract_paragraphs(parse_result)
        assert "Line one" in paras
        assert "Line two" in paras
        assert "Line three" in paras

    def test_docx_parse_result(self):
        parse_result = {
            "file_type": "docx",
            "paragraphs": [
                {"index": 0, "text": "First clause."},
                {"index": 1, "text": "Second clause."},
            ],
            "full_text": "First clause.\n\nSecond clause.",
        }
        paras = self.parser.extract_paragraphs(parse_result)
        assert paras == ["First clause.", "Second clause."]

    def test_empty_paragraphs_excluded(self):
        parse_result = {
            "file_type": "docx",
            "paragraphs": [
                {"index": 0, "text": "  "},
                {"index": 1, "text": "Real content."},
            ],
            "full_text": "Real content.",
        }
        paras = self.parser.extract_paragraphs(parse_result)
        assert paras == ["Real content."]


# ═════════════════════════════════════════════════════════════════════════════
# 3. DocumentProcessorService.process — async unit tests with mocks
# ═════════════════════════════════════════════════════════════════════════════

@pytest.mark.asyncio
class TestDocumentProcessorService:
    """Process pipeline sets correct statuses and persists data."""

    def _make_service(self, mock_db, mock_parser):
        from app.services.document_processor import DocumentProcessorService
        return DocumentProcessorService(db_service=mock_db, parser_agent=mock_parser)

    def _make_mock_db(self, existing_text=None):
        db = MagicMock()
        db.get_document_text = AsyncMock(return_value=existing_text)
        db.update_contract_status = AsyncMock()
        db.save_document_text = AsyncMock()
        db.save_chunks = AsyncMock()
        db.update_contract_page_count = AsyncMock()
        return db

    def _make_mock_parser(self, full_text="Sample contract text.", page_count=1):
        parser = MagicMock()
        parser.parse_document = AsyncMock(return_value={
            "full_text": full_text,
            "file_type": "pdf",
            "page_count": page_count,
            "pages": [{"page_number": 1, "text": full_text}],
        })
        parser.extract_paragraphs = MagicMock(return_value=["Sample contract text."])
        parser.chunk_text = MagicMock(return_value=[
            {"chunk_index": 0, "chunk_text": "Sample contract text."}
        ])
        return parser

    async def test_successful_extraction_sets_extracted_status(self):
        from app.models.schemas import ContractStatus
        db = self._make_mock_db()
        parser = self._make_mock_parser()
        svc = self._make_service(db, parser)

        await svc.process("cid-1", "/fake/path/contract.pdf", "pdf")

        calls = [c.args[1] for c in db.update_contract_status.call_args_list]
        assert ContractStatus.EXTRACTING in calls
        assert ContractStatus.EXTRACTED in calls

    async def test_successful_extraction_saves_text_and_chunks(self):
        db = self._make_mock_db()
        parser = self._make_mock_parser()
        svc = self._make_service(db, parser)

        await svc.process("cid-2", "/fake/path/contract.pdf", "pdf")

        db.save_document_text.assert_called_once()
        db.save_chunks.assert_called_once()

    async def test_empty_document_sets_failed_status(self):
        from app.models.schemas import ContractStatus
        db = self._make_mock_db()
        parser = self._make_mock_parser(full_text="   ")  # whitespace only
        svc = self._make_service(db, parser)

        await svc.process("cid-3", "/fake/path/empty.pdf", "pdf")

        final_status = db.update_contract_status.call_args_list[-1].args[1]
        assert final_status == ContractStatus.FAILED

    async def test_already_extracted_is_skipped(self):
        """If text already exists, process() must be idempotent."""
        db = self._make_mock_db(existing_text={"raw_text": "cached"})
        parser = self._make_mock_parser()
        svc = self._make_service(db, parser)

        await svc.process("cid-4", "/fake/path/contract.pdf", "pdf")

        parser.parse_document.assert_not_called()
        db.save_document_text.assert_not_called()

    async def test_parser_failure_sets_failed_status(self):
        from app.models.schemas import ContractStatus
        db = self._make_mock_db()
        parser = MagicMock()
        parser.parse_document = AsyncMock(side_effect=RuntimeError("PDF corrupted"))
        svc = self._make_service(db, parser)

        await svc.process("cid-5", "/fake/path/bad.pdf", "pdf")

        final_status = db.update_contract_status.call_args_list[-1].args[1]
        assert final_status == ContractStatus.FAILED


# ═════════════════════════════════════════════════════════════════════════════
# 4. Upload endpoint — validation tests (TestClient, patched storage + DB)
# ═════════════════════════════════════════════════════════════════════════════

@pytest.fixture(scope="module")
def client(tmp_path_factory):
    """
    Isolated TestClient with:
      - temporary SQLite DB
      - storage upload stubbed to a fake path
      - document_processor.process stubbed as no-op
    """
    tmp = tmp_path_factory.mktemp("api_test")
    db_path = str(tmp / "test.db")

    from app.services.sqlite_service import DatabaseService
    test_db = DatabaseService(db_path=db_path)

    from app.services.storage_quota_service import storage_quota_service as _quota_svc
    _original_db = _quota_svc._db
    _quota_svc._db = test_db

    from main import app
    from fastapi.testclient import TestClient

    with (
        patch("app.api.contracts.db_service", test_db),
        patch(
            "app.services.storage_factory.storage.upload_file",
            new_callable=AsyncMock,
            return_value="/fake/storage/contract.pdf",
        ),
        patch(
            "app.api.contracts.document_processor.process",
            new_callable=AsyncMock,
        ),
    ):
        yield TestClient(app)

    _quota_svc._db = _original_db


class TestUploadValidation:
    """Upload endpoint rejects invalid inputs with appropriate HTTP status codes."""

    def test_unsupported_file_type_returns_400(self, client):
        response = client.post(
            "/api/contracts/upload",
            files={"file": ("malware.exe", b"binary content", "application/octet-stream")},
        )
        assert response.status_code == 400
        assert "Unsupported file type" in response.json()["detail"]

    def test_empty_file_returns_400(self, client):
        response = client.post(
            "/api/contracts/upload",
            files={"file": ("empty.pdf", b"", "application/pdf")},
        )
        assert response.status_code == 400
        assert "empty" in response.json()["detail"].lower()

    def test_file_exceeding_size_limit_returns_400(self, client):
        # Build a file just over 50 MB (the configured limit)
        oversized = b"X" * (51 * 1024 * 1024)
        response = client.post(
            "/api/contracts/upload",
            files={"file": ("big.pdf", oversized, "application/pdf")},
        )
        assert response.status_code == 400
        assert "exceeds" in response.json()["detail"].lower()

    def test_valid_pdf_upload_returns_200(self, client):
        pdf_bytes = _make_minimal_pdf_bytes()
        response = client.post(
            "/api/contracts/upload",
            files={"file": ("contract.pdf", pdf_bytes, "application/pdf")},
        )
        assert response.status_code == 200
        body = response.json()
        assert "contract_id" in body
        assert body["status"] == "uploaded"
        assert body["filename"] == "contract.pdf"

    def test_valid_docx_upload_returns_200(self, client):
        docx_bytes = _make_docx_bytes()
        response = client.post(
            "/api/contracts/upload",
            files={"file": ("agreement.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        body = response.json()
        assert "contract_id" in body
        assert body["status"] == "uploaded"

    def test_upload_creates_db_record(self, client):
        """After a successful upload a record with status 'uploaded' must exist in DB."""
        pdf_bytes = _make_minimal_pdf_bytes()
        response = client.post(
            "/api/contracts/upload",
            files={"file": ("record_test.pdf", pdf_bytes, "application/pdf")},
        )
        assert response.status_code == 200
        contract_id = response.json()["contract_id"]

        # Retrieve the same record via GET
        get_response = client.get(
            f"/api/contracts/{contract_id}",
            params={"user_id": "anonymous"},
        )
        assert get_response.status_code == 200
        assert get_response.json()["id"] == contract_id

    def test_text_endpoint_returns_404_before_extraction(self, client):
        """GET /contracts/{id}/text must return 404 when extraction has not run."""
        pdf_bytes = _make_minimal_pdf_bytes()
        upload = client.post(
            "/api/contracts/upload",
            files={"file": ("not_extracted.pdf", pdf_bytes, "application/pdf")},
        )
        assert upload.status_code == 200
        contract_id = upload.json()["contract_id"]

        text_response = client.get(
            f"/api/contracts/{contract_id}/text",
            params={"user_id": "anonymous"},
        )
        # document_processor.process is stubbed as no-op so text is never saved
        assert text_response.status_code == 404

    def test_chunks_endpoint_returns_404_before_extraction(self, client):
        """GET /contracts/{id}/chunks must return 404 when extraction has not run."""
        pdf_bytes = _make_minimal_pdf_bytes()
        upload = client.post(
            "/api/contracts/upload",
            files={"file": ("no_chunks.pdf", pdf_bytes, "application/pdf")},
        )
        assert upload.status_code == 200
        contract_id = upload.json()["contract_id"]

        chunks_response = client.get(
            f"/api/contracts/{contract_id}/chunks",
            params={"user_id": "anonymous"},
        )
        assert chunks_response.status_code == 404
